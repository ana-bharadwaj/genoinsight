import pymongo
from docx import Document

#Connect to your MongoDB instance
client = pymongo.MongoClient("mongodb://localhost:27017")  # Replace with your MongoDB connection string

#List all the collections in the database
database = client["test7"]  # Replace with your database name
collections = database.list_collection_names()

#Sort the collections in ascending order
collections.sort()

#Create a Word document
doc = Document()

#Iterate through the collections one by one
for collection_name in collections:
    collection = database[collection_name]

    # Add the collection name as a heading
    doc.add_heading(f"Collection: {collection_name}", level=1)

    # Create a table with three columns: Document Name, Position, and End
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    table.allow_autofit = False

    # Set the table headers
    table.cell(0, 0).text = "Document Name"
    table.cell(0, 1).text = "Position"
    table.cell(0, 2).text = "End"

    # Iterate through the documents in the collection
    for i, document in enumerate(collection.find(), start=1):
        # Check if the document has a "position" field
        if "Position" in document:
            position = document["Position"]
        else:
            position = ""

        # Check if the document has an "info" field with an "end" field
        if "INFO" in document and "END" in document["INFO"]:
            end = document["INFO"]["END"]
        else:
            end = ""

        # Add a new row to the table for the current document
        table.add_row().cells[0].text = f"Document {i}"
        table.rows[i].cells[1].text = str(position)
        table.rows[i].cells[2].text = str(end)

#Save the Word document
doc.save("mongodb_datanew.docx")

#Close the MongoDB connection when done
client.close()